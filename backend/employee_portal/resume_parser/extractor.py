import os
import subprocess
from PyPDF2 import PdfReader
from docx import Document


def convert_doc_to_docx(file_path):
    """
    Convert old .doc file to .docx using LibreOffice (must be installed).
    """
    output_dir = os.path.dirname(file_path)
    subprocess.run([
        "soffice",
        "--headless",
        "--convert-to",
        "docx",
        file_path,
        "--outdir",
        output_dir
    ], check=True)

    new_file = os.path.splitext(file_path)[0] + ".docx"
    return new_file


def extract_text_from_resume(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    text = ""

    # PDF
    if ext == ".pdf":
        reader = PdfReader(file_path)
        for page in reader.pages:
            text += page.extract_text() or ""

    # Modern Word formats
    elif ext in [".docx", ".docm", ".dotx", ".dotm"]:
        doc = Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"

    # Old Word format (.doc)
    elif ext == ".doc":
        converted_file = convert_doc_to_docx(file_path)
        doc = Document(converted_file)
        for para in doc.paragraphs:
            text += para.text + "\n"

    return text.lower()




# import os
# from PyPDF2 import PdfReader
# from docx import Document


# def extract_text_from_resume(file_path):
#     ext = os.path.splitext(file_path)[1].lower()

#     text = ""

#     if ext == ".pdf":
#         reader = PdfReader(file_path)
#         for page in reader.pages:
#             text += page.extract_text() or ""

#     elif ext in [".doc", ".docx"]:
#         doc = Document(file_path)
#         for para in doc.paragraphs:
#             text += para.text + "\n"

#     return text.lower()
