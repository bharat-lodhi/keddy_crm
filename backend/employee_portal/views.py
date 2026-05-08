from rest_framework.views import APIView
from rest_framework.response import Response

from rest_framework.permissions import IsAuthenticated
from rest_framework_simplejwt.authentication import JWTAuthentication
from rest_framework.views import APIView
from rest_framework import generics

from landing.views import CsrfExemptSessionAuthentication
from .serializers import ClientSerializer,VendorDetailSerializer,VendorSingleDetailSerializer,VendorUpdateSerializer,VendorCreateSerializer,CandidateCreateSerializer,CandidateCreateSerializer,CandidateUpdateSerializer,CandidateDetailSerializer
from django.contrib.auth import get_user_model
from django.shortcuts import get_object_or_404
import os
from django.core.files.storage import default_storage
from jd_mapping.models import Requirement


from .resume_parser.extractor import extract_text_from_resume
from .resume_parser.parser import extract_skills_from_text
from .models import Candidate,Vendor,Client
from rest_framework import status
from rest_framework.generics import ListAPIView
from django.db.models import Q

User = get_user_model()

#====================verders============================================
from rest_framework.parsers import MultiPartParser, FormParser
class VendorCreateAPIView(APIView):
    parser_classes = [MultiPartParser, FormParser]
    # authentication_classes = (CsrfExemptSessionAuthentication,)
    authentication_classes = (JWTAuthentication,)
    permission_classes = (IsAuthenticated,)
    
    def post(self, request):
        user = request.user

        if user.role not in ["EMPLOYEE", "SUB_ADMIN"]:
            return Response(
                {"error": "Only employees and sub-admins can create vendors."},
                status=status.HTTP_403_FORBIDDEN
            )

        serializer = VendorCreateSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)

        vendor = serializer.save(
            uploaded_by=user,
            created_by=user,
        )

        return Response(
            {
                "message": "Vendor created successfully",
                "vendor_id": vendor.id
            },
            status=status.HTTP_201_CREATED
        )

class VendorUpdateAPIView(APIView):
    # authentication_classes = (CsrfExemptSessionAuthentication,)
    authentication_classes = (JWTAuthentication,)
    permission_classes = (IsAuthenticated,)

    def put(self, request, vendor_id):
        user = request.user

        if user.role != "EMPLOYEE":
            return Response(
                {"error": "Only employees can update vendors."},
                status=status.HTTP_403_FORBIDDEN
            )

        vendor = get_object_or_404(
            Vendor,
            id=vendor_id,
            created_by=user,
            is_deleted=False
        )

        serializer = VendorUpdateSerializer(
            vendor,
            data=request.data,
            partial=True   # Allows single field update
        )

        serializer.is_valid(raise_exception=True)
        serializer.save()

        return Response(
            {"message": "Vendor updated successfully"},
            status=status.HTTP_200_OK
        )

import logging

logger = logging.getLogger(__name__)


class IsSubAdmin:
    """
    Custom permission class for SUB_ADMIN role only
    """
    def has_permission(self, request, view):
        return request.user.is_authenticated and getattr(request.user, 'role', None) == 'SUB_ADMIN'


class VendorToggleVerifyAPIView(APIView):
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated, IsSubAdmin]

    def post(self, request, vendor_id):
        """
        Toggle vendor is_verified status (SUB_ADMIN only)
        """
        user = request.user

        # Validate vendor_id
        try:
            vendor_id = int(vendor_id)
        except (ValueError, TypeError):
            return Response(
                {
                    "success": False,
                    "error": "INVALID_VENDOR_ID",
                    "message": "Vendor ID must be a valid integer"
                },
                status=status.HTTP_400_BAD_REQUEST
            )

        # Get vendor (only not deleted)
        vendor = get_object_or_404(Vendor, id=vendor_id, is_deleted=False)

        # Toggle verification status
        old_status = vendor.is_verified
        vendor.is_verified = not vendor.is_verified
        vendor.save()

        # Log action
        logger.info(
            f"SUB_ADMIN {user.email} (ID: {user.id}) | "
            f"Vendor {vendor.company_name} (ID: {vendor.id}) | "
            f"Status changed: {old_status} -> {vendor.is_verified}"
        )

        return Response(
            {
                "success": True,
                "message": f"Vendor { 'verified' if vendor.is_verified else 'unverified' } successfully",
                "data": {
                    "vendor_id": vendor.id,
                    "vendor_name": vendor.company_name,
                    "is_verified": vendor.is_verified
                }
            },
            status=status.HTTP_200_OK
        )
        
    
class VendorCompanyPoolAPIView(ListAPIView):
    # authentication_classes = (CsrfExemptSessionAuthentication,)
    authentication_classes = (JWTAuthentication,)
    permission_classes = (IsAuthenticated,)
    serializer_class = VendorDetailSerializer

    def get_queryset(self):
        user = self.request.user

        # 🔹 Find Company Owner (SubAdmin)
        if user.role == "SUB_ADMIN":
            company_admin = user
        elif user.role == "EMPLOYEE":
            company_admin = user.parent_user
        else:
            return Vendor.objects.none()

        # 🔹 Get company users (SubAdmin + all employees)
        company_users = User.objects.filter(
            Q(id=company_admin.id) |
            Q(parent_user=company_admin)
        )

        # 🔹 Get unique vendors of company
        queryset = Vendor.objects.filter(
            created_by__in=company_users,
            is_deleted=False
        ).distinct().order_by("-created_at")

        return queryset
        
    
class UserVendorFullListAPIView(ListAPIView):
    authentication_classes = (JWTAuthentication,)
    permission_classes = (IsAuthenticated,)
    serializer_class = VendorDetailSerializer

    def get_queryset(self):
        user = self.request.user

        # Only employee allowed
        if user.role != "EMPLOYEE":
            return Vendor.objects.none()

        queryset = Vendor.objects.filter(
            Q(created_by=user) | Q(assigned_employees=user),
            is_deleted=False
        ).distinct().order_by("-created_at")

        # 🔍 GLOBAL SEARCH
        search = self.request.query_params.get("search")
        if search:
            queryset = queryset.filter(
                Q(name__icontains=search) |
                Q(company_name__icontains=search) |
                Q(email__icontains=search) |
                Q(number__icontains=search) |
                Q(company_pan_or_reg_no__icontains=search) |
                Q(poc1_name__icontains=search) |
                Q(poc2_name__icontains=search) |
                Q(top_3_clients__icontains=search) |
                Q(specialized_tech_developers__icontains=search) |
                Q(onsite_location__icontains=search)
            )

        return queryset
    
class VendorDetailAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, vendor_id):
        user = request.user

        # 🔹 Identify company admin
        if user.role == "SUB_ADMIN":
            company_admin = user
        elif user.role == "EMPLOYEE":
            company_admin = user.parent_user
        else:
            return Response({"error": "Unauthorized"}, status=403)

        # 🔹 Get all company users (SubAdmin + employees)
        company_users = User.objects.filter(
            Q(id=company_admin.id) |
            Q(parent_user=company_admin)
        )

        # 🔹 Strict company isolation
        vendor = get_object_or_404(
            Vendor,
            id=vendor_id,
            created_by__in=company_users,
            is_deleted=False
        )

        serializer = VendorSingleDetailSerializer(
            vendor,
            context={"request": request}
        )

        return Response(serializer.data)
    
class VendorSoftDeleteAPIView(APIView):
    # authentication_classes = (CsrfExemptSessionAuthentication,)
    authentication_classes = (JWTAuthentication,)
    permission_classes = (IsAuthenticated,)

    def delete(self, request, vendor_id):
        user = request.user

        # Only employee allowed
        if user.role != "EMPLOYEE":
            return Response(
                {"error": "Only employees can soft delete vendors."},
                status=status.HTTP_403_FORBIDDEN
            )

        # Employee can delete only own created vendor
        vendor = get_object_or_404(
            Vendor,
            id=vendor_id,
            created_by=user,
            is_deleted=False
        )

        vendor.is_deleted = True
        vendor.save(update_fields=["is_deleted"])

        return Response(
            {"message": "Vendor soft deleted successfully."},
            status=status.HTTP_200_OK
        )
        
        
# ===========================Clients============================================
from rest_framework import generics

class ClientCreateAPIView(generics.CreateAPIView):
    # authentication_classes = (CsrfExemptSessionAuthentication,)
    authentication_classes = (JWTAuthentication,)
    permission_classes = (IsAuthenticated,)
    serializer_class = ClientSerializer

    def perform_create(self, serializer):
        serializer.save(created_by=self.request.user)

    def create(self, request, *args, **kwargs):
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        self.perform_create(serializer)

        return Response(
            {
                "message": "Client created successfully",
                "client_id": serializer.instance.id
            },
            status=status.HTTP_201_CREATED
        )
from .serializers import ClientUpdateSerializer
class ClientUpdateAPIView(APIView):
    # authentication_classes = (CsrfExemptSessionAuthentication,)
    authentication_classes = (JWTAuthentication,)
    permission_classes = (IsAuthenticated,)

    def patch(self, request, client_id):
        user = request.user

        client = get_object_or_404(
            Client,
            id=client_id,
            is_deleted=False
        )

        allowed = False

        if user.role == "EMPLOYEE" and client.created_by == user:
            allowed = True

        elif user.role == "SUB_ADMIN" and (
            client.created_by == user or
            (client.created_by and client.created_by.parent_user == user)
        ):
            allowed = True

        if not allowed:
            return Response(
                {"error": "You do not have permission to update this client."},
                status=403
            )

        serializer = ClientUpdateSerializer(
            client,
            data=request.data,
            partial=True
        )

        serializer.is_valid(raise_exception=True)
        serializer.save()

        return Response({"message": "Client updated successfully."})




class ClientToggleVerifyAPIView(APIView):
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]

    def post(self, request, client_id):
        """
        Toggle client is_verified status (SUB_ADMIN only)
        """
        user = request.user

        # Check SUB_ADMIN role
        if user.role != "SUB_ADMIN":
            return Response(
                {
                    "success": False,
                    "error": "PERMISSION_DENIED",
                    "message": "Only SUB_ADMIN can verify or unverify clients"
                },
                status=status.HTTP_403_FORBIDDEN
            )

        # Validate client_id
        try:
            client_id = int(client_id)
        except (ValueError, TypeError):
            return Response(
                {
                    "success": False,
                    "error": "INVALID_CLIENT_ID",
                    "message": "Client ID must be a valid integer"
                },
                status=status.HTTP_400_BAD_REQUEST
            )

        # Get client (only not deleted)
        client = get_object_or_404(Client, id=client_id, is_deleted=False)

        # Toggle verification status
        old_status = client.is_verified
        client.is_verified = not client.is_verified
        client.save()

        # Log action
        logger.info(
            f"SUB_ADMIN {user.email} (ID: {user.id}) | "
            f"Client {client.company_name} (ID: {client.id}) | "
            f"Status changed: {old_status} -> {client.is_verified}"
        )

        return Response(
            {
                "success": True,
                "message": f"Client {'verified' if client.is_verified else 'unverified'} successfully",
                "data": {
                    "client_id": client.id,
                    "client_name": client.client_name,
                    "company_name": client.company_name,
                    "is_verified": client.is_verified
                }
            },
            status=status.HTTP_200_OK
        )
    
    
from rest_framework import generics
from rest_framework.permissions import IsAuthenticated


# class ClientListAPIView(generics.ListAPIView):
#     serializer_class = ClientSerializer
#     permission_classes = [IsAuthenticated]

#     def get_queryset(self):
#         user = self.request.user

#         if user.role != "EMPLOYEE":
#             return Client.objects.none()

#         return Client.objects.filter(
#             Q(created_by=user) | Q(assigned_employees=user),
#             is_deleted=False
#         ).distinct().order_by('-created_at')
        
        
from django.db.models import Q
from rest_framework.filters import SearchFilter
from django_filters.rest_framework import DjangoFilterBackend
from rest_framework import generics
from rest_framework.permissions import IsAuthenticated

class ClientListAPIView(generics.ListAPIView):
    serializer_class = ClientSerializer
    permission_classes = [IsAuthenticated]

    filter_backends = [SearchFilter]
    search_fields = [
        "client_name",
        "company_name",
        "phone_number",
    ]

    def get_queryset(self):
        user = self.request.user

        if user.role != "EMPLOYEE":
            return Client.objects.none()

        return Client.objects.filter(
            Q(created_by=user) | Q(assigned_employees=user),
            is_deleted=False
        ).distinct().order_by("-created_at")
                
    
from .serializers import ClientDetailSerializer
class ClientDetailAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, client_id):
        user = request.user

        client = get_object_or_404(
            Client,
            id=client_id,
            is_deleted=False
        )

        allowed = False

        # Employee created
        if user.role == "EMPLOYEE" and client.created_by == user:
            allowed = True

        # Employee assigned
        elif user.role == "EMPLOYEE" and user in client.assigned_employees.all():
            allowed = True

        # SubAdmin own client
        elif user.role == "SUB_ADMIN" and client.created_by == user:
            allowed = True

        # SubAdmin employee client
        elif (
            user.role == "SUB_ADMIN"
            and client.created_by
            and client.created_by.parent_user == user
        ):
            allowed = True

        if not allowed:
            return Response(
                {"error": "You do not have permission to view this client."},
                status=403
            )

        serializer = ClientDetailSerializer(
            client,
            context={"request": request}
        )

        return Response(serializer.data)

class ClientSoftDeleteAPIView(APIView):
    # authentication_classes = (CsrfExemptSessionAuthentication,)
    authentication_classes = (JWTAuthentication,)
    permission_classes = (IsAuthenticated,)

    def delete(self, request, client_id):
        user = request.user

        if user.role != "EMPLOYEE":
            return Response(
                {"error": "Only employees can soft delete clients."},
                status=403
            )

        client = get_object_or_404(
            Client,
            id=client_id,
            created_by=user,
            is_deleted=False
        )

        client.is_deleted = True
        client.save(update_fields=["is_deleted"])

        return Response(
            {"message": "Client soft deleted successfully."},
            status=200
        )

# =========================Candidate==========================================


from rest_framework.views import APIView
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from django.db.models import Q
from landing.models import User


class EmployeeDropdownAPIView(APIView):
    permission_classes = (IsAuthenticated,)

    def get(self, request):
        user = request.user
        search_query = request.query_params.get('search', '').strip()

        # Identify company admin
        if user.role == "SUB_ADMIN":
            company_admin = user
        elif user.role == "EMPLOYEE":
            company_admin = user.parent_user
        else:
            return Response([], status=200)

        # Base queryset
        employees = User.objects.filter(
            parent_user=company_admin,
            role="EMPLOYEE"
        )

        # Search filter - id, first_name, last_name
        if search_query:
            employees = employees.filter(
                Q(id__icontains=search_query) |
                Q(first_name__icontains=search_query) |
                Q(last_name__icontains=search_query)
            )

        # Same response format - no change
        employees_data = employees.values("id", "first_name", "last_name")

        return Response(employees_data)
    
# -------------------------

import io
from pdfminer.high_level import extract_text
from docx import Document

from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from rest_framework.permissions import IsAuthenticated

from .serializers import ResumeUploadSerializer
from .utils import parse_resume_text


# def read_resume_file(uploaded_file):
#     filename = uploaded_file.name.lower()

#     if filename.endswith(".pdf"):
#         file_stream = io.BytesIO(uploaded_file.read())
#         return extract_text(file_stream)

#     elif filename.endswith(".docx"):
#         document = Document(uploaded_file)
#         return "\n".join([p.text for p in document.paragraphs])

#     elif filename.endswith(".txt"):
#         return uploaded_file.read().decode("utf-8")

#     else:
#         raise ValueError("Unsupported file format")


# class ResumeParseAPIView(APIView):
#     authentication_classes = (CsrfExemptSessionAuthentication,)
#     permission_classes = (IsAuthenticated,)

#     def post(self, request):
#         serializer = ResumeUploadSerializer(data=request.data)
#         serializer.is_valid(raise_exception=True)

#         resume_file = serializer.validated_data['resume']

#         text = read_resume_file(resume_file)

#         parsed_data = parse_resume_text(text)

#         return Response(
#             {
#                 "message": "Resume parsed successfully",
#                 "data": parsed_data
#             },
#             status=status.HTTP_200_OK
#         )


import io
import os
import subprocess
from PyPDF2 import PdfReader
from docx import Document


def convert_doc_to_docx_temp(uploaded_file):
    """
    Convert old .doc file to .docx using LibreOffice.
    """
    import tempfile

    with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as temp_doc:
        temp_doc.write(uploaded_file.read())
        temp_doc_path = temp_doc.name

    output_dir = os.path.dirname(temp_doc_path)

    subprocess.run([
        "soffice",
        "--headless",
        "--convert-to",
        "docx",
        temp_doc_path,
        "--outdir",
        output_dir
    ], check=True)

    new_file_path = os.path.splitext(temp_doc_path)[0] + ".docx"
    return new_file_path


def read_resume_file(uploaded_file):
    filename = uploaded_file.name.lower()

    # PDF
    if filename.endswith(".pdf"):
        file_stream = io.BytesIO(uploaded_file.read())
        reader = PdfReader(file_stream)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text

    # Modern Word formats
    elif filename.endswith((".docx", ".docm", ".dotx", ".dotm")):
        document = Document(uploaded_file)
        return "\n".join([p.text for p in document.paragraphs])

    # Old Word format (.doc)
    elif filename.endswith(".doc"):
        converted_path = convert_doc_to_docx_temp(uploaded_file)
        document = Document(converted_path)
        return "\n".join([p.text for p in document.paragraphs])

    # TXT
    elif filename.endswith(".txt"):
        return uploaded_file.read().decode("utf-8")

    else:
        raise ValueError("Unsupported file format")


class ResumeParseAPIView(APIView):
    # authentication_classes = (CsrfExemptSessionAuthentication,)
    authentication_classes = (JWTAuthentication,)
    permission_classes = (IsAuthenticated,)

    def post(self, request):
        serializer = ResumeUploadSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)

        resume_file = serializer.validated_data['resume']

        text = read_resume_file(resume_file)

        parsed_data = parse_resume_text(text)

        return Response(
            {
                "message": "Resume parsed successfully",
                "data": parsed_data
            },
            status=status.HTTP_200_OK
        )


#=========================Candidate=========================


class CandidateCreateAPIView(generics.CreateAPIView):
    authentication_classes = (JWTAuthentication,)
    permission_classes = [IsAuthenticated,]
    queryset = Candidate.objects.all()
    serializer_class = CandidateCreateSerializer
    permission_classes = (IsAuthenticated,)

    def create(self, request, *args, **kwargs):
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        self.perform_create(serializer)

        return Response(
            {
                "message": "Candidate created successfully",
                "data": serializer.data
            },
            status=status.HTTP_201_CREATED
        )


from rest_framework import generics
from rest_framework.permissions import IsAuthenticated
from django_filters.rest_framework import DjangoFilterBackend
from rest_framework.filters import OrderingFilter, SearchFilter

from .serializers import CandidateListSerializer
from .candidate_filters import CandidateFilter


class CandidateListAPIView(generics.ListAPIView):
    serializer_class = CandidateListSerializer
    permission_classes = (IsAuthenticated,)
    filter_backends = [DjangoFilterBackend, OrderingFilter, SearchFilter]
    filterset_class = CandidateFilter

    search_fields = [
        "candidate_name",
        "candidate_email",
        "skills",
        "technology",
        "vendor__company_name",   # Updated (live relation search)
    ]

    ordering_fields = "__all__"

    def get_queryset(self):
        user = self.request.user

        # Identify company admin
        if user.role == "SUB_ADMIN":
            company_admin = user
        elif user.role == "EMPLOYEE":
            company_admin = user.parent_user
        else:
            return Candidate.objects.none()

        # Company users (SubAdmin + employees)
        company_users = User.objects.filter(
            Q(id=company_admin.id) | Q(parent_user=company_admin)
        )

        return Candidate.objects.filter(
            created_by__in=company_users
        ).select_related(
            "vendor",
            "created_by",
            "submitted_to"
        ).order_by("-created_at")
        


from .serializers import CandidateListSerializer

from django.db.models import Q
from django.contrib.auth import get_user_model
from rest_framework import generics
from rest_framework.permissions import IsAuthenticated
from django_filters.rest_framework import DjangoFilterBackend
from rest_framework.filters import OrderingFilter, SearchFilter


class SubmittedProfilesView(generics.ListAPIView):
    serializer_class = CandidateListSerializer
    permission_classes = [IsAuthenticated]

    filter_backends = [DjangoFilterBackend, OrderingFilter, SearchFilter]
    filterset_class = CandidateFilter

    search_fields = [
        "candidate_name",
        "candidate_email",
        "skills",
        "technology",
        "vendor__company_name",
    ]

    ordering_fields = "__all__"

    def get_company_root(self, user):
        if user.role == "SUB_ADMIN":
            return user
        if user.role == "EMPLOYEE" and user.parent_user:
            return user.parent_user
        return None

    def get_queryset(self):
        user = self.request.user
        UserModel = get_user_model()

        company_root = self.get_company_root(user)
        if not company_root:
            return Candidate.objects.none()

        company_users = UserModel.objects.filter(
            Q(id=company_root.id) | Q(parent_user=company_root)
        )

        return (
            Candidate.objects.select_related(
                "vendor",
                "client",
                "created_by",
                "submitted_to",
            )
            .filter(
                Q(created_by=user) | Q(submitted_to=user),
                verification_status=True,
                client__isnull=False,
                created_by__in=company_users,
                is_deleted = False
            )
            .order_by("-created_at")
        )


from django.db.models import Q
from django.contrib.auth import get_user_model
from rest_framework import generics, status
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework.exceptions import PermissionDenied


class CandidateSoftDeleteAPIView(generics.DestroyAPIView):
    permission_classes = (IsAuthenticated,)

    def get_company_root(self, user):
        if user.role == "SUB_ADMIN":
            return user
        if user.role == "EMPLOYEE" and user.parent_user:
            return user.parent_user
        raise PermissionDenied("Invalid company structure.")

    def get_queryset(self):
        user = self.request.user
        UserModel = get_user_model()

        company_root = self.get_company_root(user)

        company_users = UserModel.objects.filter(
            Q(id=company_root.id) | Q(parent_user=company_root)
        )

        return Candidate.objects.filter(
            created_by__in=company_users,
            is_deleted=False
        )

    def destroy(self, request, *args, **kwargs):
        instance = self.get_object()
        user = request.user

        # Only creator employee or company SubAdmin can soft delete
        if user.role == "SUB_ADMIN":
            pass
        elif instance.created_by == user:
            pass
        else:
            raise PermissionDenied("You do not have permission to delete this candidate.")

        instance.is_deleted = True
        instance.changed_by = user
        instance.save(update_fields=["is_deleted", "changed_by"])

        return Response(
            {"message": "Candidate soft deleted successfully"},
            status=status.HTTP_200_OK
        )        
        

from django_filters.rest_framework import DjangoFilterBackend
from rest_framework import generics
from rest_framework.permissions import IsAuthenticated
from rest_framework.filters import OrderingFilter, SearchFilter
from rest_framework.exceptions import PermissionDenied
from django.contrib.auth import get_user_model
User = get_user_model()

class UserCandidateListAPIView(generics.ListAPIView):
    serializer_class = CandidateListSerializer
    permission_classes = (IsAuthenticated,)

    filter_backends = [DjangoFilterBackend, OrderingFilter, SearchFilter]
    filterset_class = CandidateFilter

    search_fields = [
        "candidate_name",
        "candidate_email",
        "skills",
        "technology",
        "vendor__company_name",
    ]

    ordering_fields = "__all__"

    def get_company_root(self, user):
        if user.role == "SUB_ADMIN":
            return user
        if user.role == "EMPLOYEE" and user.parent_user:
            return user.parent_user
        raise PermissionDenied("Invalid company structure.")

    def get_queryset(self):
        user = self.request.user
        company_root = self.get_company_root(user)

        company_users = User.objects.filter(
            Q(id=company_root.id) | Q(parent_user=company_root)
        )

        return (
            Candidate.objects.select_related(
                "vendor",
                "client",
                "created_by",
                "submitted_to",
            )
            .filter(
                Q(created_by=user) |
                Q(submitted_to=user, client__isnull=False),
                created_by__in=company_users,
                vendor__created_by__in=company_users,
                is_deleted = False
            )
            .order_by("-created_at")
        )

# from django.db.models import Q
# from rest_framework.decorators import api_view, permission_classes
# from rest_framework.permissions import IsAuthenticated
# from rest_framework.response import Response

# # Submitted profile new navbar option
# @api_view(["GET"])
# @permission_classes([IsAuthenticated])
# def submitted_profiles(request):
#     user = request.user

#     candidates = Candidate.objects.filter(
#         Q(created_by=user) | Q(submitted_to=user),
#         verification_status=True,
#         client__isnull=False
#     ).select_related("vendor", "client").order_by("-created_at")

#     serializer = TodayCandidateSerializer(candidates, many=True)
#     return Response(serializer.data)

from django.db.models import Q
from rest_framework import generics, status
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework.exceptions import PermissionDenied


class CandidateUpdateAPIView(generics.RetrieveUpdateAPIView):
    authentication_classes = (JWTAuthentication,)
    permission_classes = (IsAuthenticated,)
    serializer_class = CandidateUpdateSerializer

    def get_company_root(self, user):
        if user.role == "SUB_ADMIN":
            return user
        if user.role == "EMPLOYEE" and user.parent_user:
            return user.parent_user
        raise PermissionDenied("Invalid company structure.")

    def get_queryset(self):
        user = self.request.user
        company_root = self.get_company_root(user)

        from django.contrib.auth import get_user_model

        UserModel = get_user_model()

        company_users = UserModel.objects.filter(
            Q(id=company_root.id) | Q(parent_user=company_root)
        )

        return Candidate.objects.select_related(
            "vendor",
            "client",
            "created_by",
            "submitted_to",
        ).filter(
            created_by__in=company_users
        )

    def update(self, request, *args, **kwargs):
        partial = True
        instance = self.get_object()
        user = request.user

        # Permission Logic
        if user.role == "SUB_ADMIN":
            pass
        elif instance.created_by == user:
            pass
        elif instance.submitted_to == user:
            pass
        else:
            raise PermissionDenied("You do not have permission to update this candidate.")

        serializer = self.get_serializer(
            instance,
            data=request.data,
            partial=partial,
            context={"request": request}
        )
        serializer.is_valid(raise_exception=True)

        serializer.save(changed_by=user)

        return Response(
            {
                "message": "Candidate updated successfully",
                "data": serializer.data
            },
            status=status.HTTP_200_OK
        )
        
        
from rest_framework import generics
from rest_framework.permissions import IsAuthenticated


class CandidateDetailAPIView(generics.RetrieveAPIView):
    serializer_class = CandidateDetailSerializer
    permission_classes = (IsAuthenticated,)

    def get_company_root(self, user):
        if user.role == "SUB_ADMIN":
            return user
        if user.role == "EMPLOYEE" and user.parent_user:
            return user.parent_user
        raise PermissionDenied("Invalid company structure.")

    def get_queryset(self):
        user = self.request.user
        UserModel = get_user_model()

        company_root = self.get_company_root(user)

        company_users = UserModel.objects.filter(
            Q(id=company_root.id) | Q(parent_user=company_root)
        )

        return Candidate.objects.select_related(
            "vendor",
            "client",
            "created_by",
            "submitted_to",
        ).filter(
            created_by__in=company_users
        )


#=============================Employee Dashboard====================================================

from django.utils.timezone import now
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response

from employee_portal.models import Vendor, Client, Candidate
from .serializers import DashboardStatsSerializer

from django.db.models import Q
from django.utils.timezone import now
from datetime import timedelta

from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response


@api_view(["GET"])
@permission_classes([IsAuthenticated])
def dashboard_stats(request):
    user = request.user

    # 🚫 Block SubAdmin access
    if user.role != "EMPLOYEE":
        return Response({"detail": "Only employees allowed."}, status=403)

    today = now().date()
    two_days_ago = now() - timedelta(days=2)

    # ===== Vendors (exclude soft deleted) =====
    total_vendors = Vendor.objects.filter(
        Q(created_by=user) | Q(assigned_employees=user),
        is_deleted=False
    ).count()

    # ===== Clients (exclude soft deleted) =====
    total_clients =Client.objects.filter(
        Q(created_by=user) | Q(assigned_employees=user),
        is_deleted=False
            ).distinct().count()

    # ===== Candidates (exclude soft deleted) =====
    total_profiles = Candidate.objects.filter(
        created_by=user,
        is_deleted=False
    ).count()

    today_profiles = Candidate.objects.filter(
        created_by=user,
        created_at__date=today,
        is_deleted=False
    ).count()

    # ===== Today Submitted Profiles =====
    today_submitted_profiles = Candidate.objects.filter(
        created_at__date=today,
        verification_status=True,
        is_deleted=False,
        client__isnull=False
    ).filter(
        Q(created_by=user) |
        Q(submitted_to=user)
    ).count()

    # ===== Pipeline Count =====
    total_pipelines = Candidate.objects.filter(
        Q(created_by=user) | Q(submitted_to=user),
        verification_status=True,
        is_deleted=False
    ).filter(
        Q(main_status__iexact="SCREENING") |
        Q(main_status__iexact="L1") |
        Q(main_status__iexact="L2") |
        Q(main_status__iexact="L3") |
        Q(main_status__iexact="OTHER")
    ).exclude(
        sub_status="REJECTED"
    ).filter(
        Q(sub_status="ON_HOLD", created_at__gte=two_days_ago) |
        ~Q(sub_status="ON_HOLD")
    ).count()
    
    # ===== Today's Requirements (JDs) =====
    # JDs created by user today
    today_created_jds = Requirement.objects.filter(
        created_by=user,
        created_at__date=today,
        is_deleted=False
    ).count()
    
    # JDs assigned to user today
    today_assigned_jds = Requirement.objects.filter(
        assignments__assigned_to=user,
        created_at__date=today,
        is_deleted=False
    ).distinct().count()
    
    # Total today's requirements (created + assigned)
    today_total_requirements = today_created_jds + today_assigned_jds
    
    # ==============================================================

    data = {
        "user_name": user.get_full_name() or user.email,
        "total_vendors": total_vendors,
        "total_clients": total_clients,
        "total_profiles": total_profiles,
        "today_profiles": today_profiles,
        "today_submitted_profiles": today_submitted_profiles,
        "total_pipelines": total_pipelines,
        "today_requirements": today_total_requirements,
    }

    return Response(data)


from .serializers import TodayCandidateSerializer

from django.utils.timezone import now
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response


@api_view(["GET"])
@permission_classes([IsAuthenticated])
def today_user_candidates(request):
    user = request.user

    # 🚫 Only Employee allowed
    if user.role != "EMPLOYEE":
        return Response({"detail": "Only employees allowed."}, status=403)

    today = now().date()

    candidates = (
        Candidate.objects.filter(
            created_by=user,
            created_at__date=today,
            is_deleted=False  # ✅ Exclude soft deleted
        )
        .select_related("vendor", "client")
        .order_by("-created_at")
    )

    serializer = TodayCandidateSerializer(candidates, many=True)
    return Response(serializer.data)


from django.db.models import Q
from django.utils.timezone import now

@api_view(["GET"])
@permission_classes([IsAuthenticated])
def today_verified_candidates(request):
    user = request.user

    # 🚫 Only Employee allowed
    if user.role != "EMPLOYEE":
        return Response({"detail": "Only employees allowed."}, status=403)

    today = now().date()

    candidates = (
        Candidate.objects.filter(
            created_at__date=today,
            verification_status=True,
            is_deleted=False,  # ✅ Exclude soft deleted
            client__isnull=False
        )
        .filter(
            Q(created_by=user) |
            Q(submitted_to=user)
            
        )
        .select_related("vendor", "client")
        .order_by("-created_at")
    )

    serializer = TodayCandidateSerializer(candidates, many=True)
    return Response(serializer.data)


@api_view(["GET"])
@permission_classes([IsAuthenticated])
def last_7_days_verified_candidates(request):
    user = request.user

    # 🚫 Only Employee allowed
    if user.role != "EMPLOYEE":
        return Response({"detail": "Only employees allowed."}, status=403)

    today = now().date()
    seven_days_ago = today - timedelta(days=7)

    candidates = (
        Candidate.objects.filter(
            created_at__date__range=(seven_days_ago, today),
            verification_status=True,
            is_deleted=False  # ✅ Exclude soft deleted
        )
        .filter(
            Q(created_by=user) |
            Q(submitted_to=user, client__isnull=False)
        )
        .select_related("vendor", "client")
        .order_by("-created_at")
    )

    serializer = TodayCandidateSerializer(candidates, many=True)
    return Response(serializer.data)

from django.utils.timezone import now
from datetime import timedelta

@api_view(["GET"])
@permission_classes([IsAuthenticated])
def active_pipeline_candidates(request):
    user = request.user

    # 🚫 Only Employee allowed
    if user.role != "EMPLOYEE":
        return Response({"detail": "Only employees allowed."}, status=403)

    two_days_ago = now() - timedelta(days=2)

    candidates = (
        Candidate.objects.filter(
            Q(created_by=user) | Q(submitted_to=user),
            verification_status=True,
            is_deleted=False  # ✅ Exclude soft deleted
        )
        .filter(
            Q(main_status__iexact="SCREENING") |
            Q(main_status__iexact="L1") |
            Q(main_status__iexact="L2") |
            Q(main_status__iexact="L3") |
            Q(main_status__iexact="OTHER")
        )
        .exclude(
            sub_status="REJECTED"
        )
        .filter(
            Q(sub_status="ON_HOLD", created_at__gte=two_days_ago) |
            ~Q(sub_status="ON_HOLD")
        )
        .select_related("vendor", "client")
        .order_by("-created_at")
    )

    serializer = TodayCandidateSerializer(candidates, many=True)
    return Response(serializer.data)

from django.utils.timezone import now
from datetime import timedelta
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response


@api_view(["GET"])
@permission_classes([IsAuthenticated])
def today_team_submissions(request):
    user = request.user

    # 🚫 Only Employee allowed
    if user.role != "EMPLOYEE":
        return Response({"detail": "Only employees allowed."}, status=403)

    today = now().date()

    candidates = (
        Candidate.objects.filter(
            submitted_to=user,
            created_at__date=today,
            is_deleted=False  # ✅ Exclude soft deleted
        )
        .exclude(
            created_by=user
        )
        .select_related("vendor", "client")
        .order_by("-created_at")
    )

    serializer = TodayCandidateSerializer(candidates, many=True)
    return Response(serializer.data)


@api_view(["GET"])
@permission_classes([IsAuthenticated])
def all_team_submissions(request):
    user = request.user

    # 🚫 Only Employee allowed
    if user.role != "EMPLOYEE":
        return Response({"detail": "Only employees allowed."}, status=403)

    candidates = (
        Candidate.objects.filter(
            submitted_to=user,
            is_deleted=False  # ✅ Exclude soft deleted
        )
        .exclude(
            created_by=user
        )
        .select_related("vendor", "client")
        .order_by("-created_at")
    )

    serializer = TodayCandidateSerializer(candidates, many=True)
    return Response(serializer.data)

